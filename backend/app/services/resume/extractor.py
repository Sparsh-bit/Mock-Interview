"""
Resume text extraction — services/resume/extractor.py

Turns an uploaded PDF or DOCX into plain text the AI analyser can read.

Pure functions over bytes: no database, no network, no filesystem. That keeps them
testable without fixtures on disk, and means extraction can run on the bytes we
already hold in memory at upload time rather than downloading the file back out of
Supabase Storage afterwards.

Deliberately conservative about failure. A resume that cannot be read must produce
a clear, catchable error, because the alternative -- storing empty text and
carrying on -- is what left every upload sitting at parsing_status="pending" while
the interview silently ignored the resume. Silence is the failure mode to avoid.

"CANNOT BE READ" IS NOT ONLY "IS EMPTY", which is what this file used to assume. A
file clears a length check and is still unusable when it is a phone scan whose only
text layer is the scanner app's own page furniture, or a PDF whose fonts carry no
usable character map and extracts to a wall of U+FFFD. Both were measured passing
straight through to the AI analyser, which found no skills in them -- correctly --
after which the upload told the candidate "your resume was read successfully, but
the detailed skill analysis could not be completed". That sentence is a dead end for
someone holding a scanned PDF. Every check below exists to replace it with the one
thing they can act on: which file to upload instead.
"""

from __future__ import annotations

import io
import re

import structlog

logger = structlog.get_logger(__name__)

#: MIME types we can actually extract, mapped to a short kind for dispatch.
_PDF_MIMES = frozenset(
    {
        "application/pdf",
        "application/x-pdf",
    }
)
_DOCX_MIMES = frozenset(
    {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/docx",
    }
)

#: Below this, extraction "succeeded" but produced nothing usable. The usual cause
#: is a scanned/photographed resume: a PDF whose pages are images, with no text
#: layer at all. That needs OCR, which we do not do -- so say so plainly instead of
#: handing the interviewer three characters of noise.
_MIN_USEFUL_CHARS = 200

#: Below this, text that shows none of the usual resume markers is treated as junk
#: rather than as an unusual resume.
#:
#: THE 200-CHAR FLOOR ABOVE IS NOT ENOUGH ON ITS OWN, and this is measured, not
#: theoretical. A phone-scanned resume is not always a page of pure images: the
#: scanner app stamps its own furniture into a text layer, and four pages of
#: "Scanned by CamScanner  Page 3 of 8  IMG_0411.jpg" is 307 characters — clear of
#: the floor, so extraction reported success and the AI analyser was handed
#: scanner branding to find skills in. Both analysis halves then correctly found
#: nothing, four billed retries later the upload stored `text_only`, and the
#: candidate was told "your resume was read successfully" about a file the
#: interviewer can make no use of. That message is worse than an error: it is
#: unactionable, and it leaves them believing their interviews are personalised.
#:
#: TWO SIGNALS TOGETHER, because either alone would reject real resumes. Length
#: alone would fail a genuinely terse one-page fresher CV. A marker check alone
#: would fail an unusual-but-substantial document — an academic CV whose headings
#: are all "Publications" and "Positions Held".
#:
#: AND THE MARKER BAR HERE IS ZERO, not the two that `looks_like_a_resume` wants.
#: That is a measured correction, not caution for its own sake: at "fewer than two"
#: this gate rejected a real 297-character fresher CV that said "B.E. Computer
#: Science, Anna University" and then just listed its projects without ever writing
#: the word "projects". One marker is weak evidence FOR a resume; zero markers in a
#: short document is strong evidence against one, and refusing a candidate's actual
#: resume is a worse failure than analysing a thin one. The measured junk this
#: catches — scanner page furniture, cover sheets, tickets — has none at all.
#:
#: 1200 characters is roughly a third of a one-page resume, so anything longer is
#: given the benefit of the doubt and passed to the analyser regardless.
_MIN_RESUME_LIKE_CHARS = 1200

#: Share of characters that may be unreadable before the text layer is declared
#: broken.
#:
#: A PDF whose font has no usable ToUnicode CMap extracts to U+FFFD replacement
#: characters or raw control codes — 560 characters of "���" clears
#: every length check there is and looks, to every downstream component, exactly
#: like a successfully read resume. This is the one unambiguous signal in this
#: file: no real resume contains replacement characters or C0 control codes, in any
#: language. Letters, digits, punctuation and every script from Devanagari to CJK
#: are all "readable" here, so a non-English resume is unaffected.
#:
#: 10%, not zero, because a single stray glyph from one bad ligature must not cost
#: a candidate an otherwise perfect resume.
_MAX_UNREADABLE_SHARE = 0.10

#: Upper bound on the text we keep. A resume is one or two pages; anything far
#: beyond that is a portfolio, a thesis, or a PDF with a pathological text layer.
#: The cap matters because this text becomes AI input on every interview plan --
#: unbounded here means unbounded token spend later.
MAX_RESUME_CHARS = 20_000


class ResumeExtractionError(Exception):
    """
    Raised when a resume's text cannot be extracted.

    Carries a message written for the candidate, not the developer: it is surfaced
    directly in the upload response so they know whether to re-export the file,
    convert it, or type their details instead.
    """

    def __init__(self, message: str, *, reason: str) -> None:
        super().__init__(message)
        #: Short machine-readable cause, for logs and metrics.
        self.reason = reason


def normalise_whitespace(text: str) -> str:
    """
    Collapse the ragged whitespace PDF extraction produces.

    PDF text layers have no concept of a paragraph: extraction yields hard line
    breaks mid-sentence, runs of spaces where the layout used columns, and blank
    lines between every bullet. Left alone this wastes tokens and makes the text
    harder for the model to read, so lines are joined and blank runs collapsed
    while genuine paragraph breaks are preserved.
    """
    # Normalise line endings first so the paragraph rules below see one form.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Non-breaking spaces and zero-width characters are common in exported PDFs
    # and would otherwise survive into the prompt as invisible noise.
    text = text.replace(" ", " ").replace("​", "")
    # Three or more newlines is never meaningful structure — cap at a blank line.
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse runs of spaces/tabs, but not newlines.
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Trim each line so indentation from a two-column layout does not survive.
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


#: The words a resume of any shape almost always contains at least one of.
_RESUME_MARKERS = (
    "experience",
    "education",
    "skill",
    "project",
    "internship",
    "certification",
    "achievement",
    "objective",
    "summary",
    "college",
    "university",
    "b.tech",
    "bachelor",
)


def resume_marker_count(text: str) -> int:
    """
    How many of the usual resume words appear. The raw signal behind the two
    judgements below, which need different amounts of it.
    """
    lowered = text.lower()
    return sum(1 for marker in _RESUME_MARKERS if marker in lowered)


def looks_like_a_resume(text: str) -> bool:
    """
    Cheap sanity check that the extracted text is plausibly a resume.

    Not validation and not security -- purely a guard against the common mistake
    of uploading the wrong PDF (a ticket, an offer letter, a question bank) and
    then wondering why the interviewer asks about nothing on it. Two or more of
    the usual section headings is enough signal; a real resume always has several.

    Used as a SOFT signal only (the upload endpoint logs it). The hard rejection in
    extract_text deliberately uses a stricter bar -- zero markers rather than fewer
    than two -- because refusing a real resume is a worse outcome than analysing a
    weak one. See _MIN_RESUME_LIKE_CHARS.
    """
    return resume_marker_count(text) >= 2


def unreadable_share(text: str) -> float:
    """
    Fraction of `text` that no resume could legitimately contain.

    Counts U+FFFD (the decoder's "I could not map this glyph" marker) and C0/C1
    control codes, excluding the tab/newline/carriage-return that real documents
    use. Everything else — every alphabet, every digit, every punctuation mark —
    counts as readable, so this cannot flag a resume for being in another language.
    """
    if not text:
        return 0.0
    unreadable = sum(
        1
        for char in text
        if char == "\ufffd" or (ord(char) < 32 and char not in "\t\n\r") or 0x7F <= ord(char) < 0xA0
    )
    return unreadable / len(text)


def _extract_pdf(data: bytes) -> str:
    """Extract text from every page of a PDF."""
    # pypdf, NOT PyPDF2. PyPDF2 is unmaintained and its last release predates several
    # parser hardening fixes — which matters here more than almost anywhere else in the
    # codebase, because this function is handed a file an anonymous user uploaded and asked
    # to parse it. pypdf is the same project's maintained successor with a compatible API.
    from pypdf import PdfReader  # noqa: PLC0415
    from pypdf.errors import PdfReadError  # noqa: PLC0415

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise ResumeExtractionError(
            "That PDF could not be opened — it may be corrupted. Try re-exporting "
            "or re-saving it, then upload again.",
            reason="pdf_unreadable",
        ) from exc

    # An encrypted PDF yields empty pages rather than raising, so check first and
    # give the real reason instead of the generic "no text found".
    if getattr(reader, "is_encrypted", False):
        # A blank owner password is common on "protected" exports and often opens.
        # decrypt() returns a PasswordType enum whose members are falsy only for
        # failure, so coerce to bool rather than comparing against an int.
        opened = False
        try:
            opened = bool(reader.decrypt(""))
        except Exception:  # noqa: BLE001 - any failure means we cannot read it
            opened = False
        if not opened:
            raise ResumeExtractionError(
                "That PDF is password-protected, so its text cannot be read. "
                "Upload an unprotected copy.",
                reason="pdf_encrypted",
            )

    parts: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page must not lose the rest
            logger.warning("resume_pdf_page_extract_failed", page=index)
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """Extract text from a DOCX: paragraphs plus table cells."""
    import docx  # noqa: PLC0415

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - python-docx raises several types
        raise ResumeExtractionError(
            "That DOCX could not be opened — it may be corrupted, or it may be an "
            "older .doc file saved with a .docx name. Re-save it as .docx and try "
            "again.",
            reason="docx_unreadable",
        ) from exc

    parts = [p.text for p in document.paragraphs]
    # Many resume templates lay everything out in a table, so paragraphs alone
    # would come back nearly empty and look like a scanned file.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text(data: bytes, mime_type: str, *, filename: str = "") -> str:
    """
    Extract normalised plain text from resume file bytes.

    Raises ResumeExtractionError with a candidate-facing message if the file
    cannot be read, is the wrong format, or contains no text layer.
    """
    if not data:
        raise ResumeExtractionError(
            "That file is empty. Upload the resume file itself, not a shortcut or "
            "link to it.",
            reason="empty_file",
        )

    kind = mime_type.split(";")[0].strip().lower()
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Browsers occasionally send a generic or wrong content type (notably
    # application/octet-stream from some file managers), so fall back to the
    # extension rather than rejecting a perfectly good PDF.
    if kind in _PDF_MIMES or (kind not in _DOCX_MIMES and suffix == "pdf"):
        raw = _extract_pdf(data)
    elif kind in _DOCX_MIMES or suffix == "docx":
        raw = _extract_docx(data)
    else:
        raise ResumeExtractionError(
            f"Unsupported resume format ({mime_type or 'unknown'}). Upload a PDF "
            "or DOCX.",
            reason="unsupported_type",
        )

    text = normalise_whitespace(raw)

    if len(text) < _MIN_USEFUL_CHARS:
        raise ResumeExtractionError(
            "No readable text was found in that file. If it is a scan or a photo "
            "of your resume, the text cannot be read — upload the original PDF "
            "exported from your editor instead.",
            reason="no_text_layer",
        )

    # ── The text exists. Is any of it usable? ───────────────────────────────
    #
    # Everything below this line is the difference between an error the candidate
    # can act on and an interview quietly conducted on noise. Both checks were
    # added after the reported bug: neither of these files was rejected, both were
    # stored as successfully-read resumes, and both produced an empty analysis that
    # the upload then explained away with "your resume was read successfully".
    share = unreadable_share(text)
    if share > _MAX_UNREADABLE_SHARE:
        logger.warning(
            "resume_text_layer_unreadable",
            chars=len(text),
            unreadable_share=round(share, 3),
        )
        raise ResumeExtractionError(
            "The text in that file could not be decoded — its fonts do not carry "
            "readable character information, which usually happens with an older "
            "or unusual PDF export. Re-export it as a PDF (or save it as DOCX) and "
            "upload it again.",
            reason="text_unreadable",
        )

    if len(text) < _MIN_RESUME_LIKE_CHARS and resume_marker_count(text) == 0:
        logger.warning(
            "resume_content_not_a_resume",
            chars=len(text),
            preview=text[:120],
        )
        raise ResumeExtractionError(
            "Only a little text could be read from that file, and it does not look "
            "like a resume — a scanner app's page markers, for instance. If you "
            "scanned or photographed your resume, upload the original PDF exported "
            "from your editor instead; otherwise check you picked the right file.",
            reason="no_resume_content",
        )

    if len(text) > MAX_RESUME_CHARS:
        logger.info(
            "resume_text_truncated",
            original_chars=len(text),
            kept_chars=MAX_RESUME_CHARS,
        )
        text = text[:MAX_RESUME_CHARS]

    return text
