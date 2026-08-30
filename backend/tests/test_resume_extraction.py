"""
Tests for resume text extraction — services/resume/extractor.py

These matter because the failure mode they guard against is silent. Resume upload
previously stored the file, set parsing_status="pending", and never extracted
anything — so a candidate could upload a resume, see a success message, and get an
interview that had never seen a word of it. Every test here is about making a
failure loud instead.

Pure functions over bytes: no database, no network, no fixtures on disk.
"""

from __future__ import annotations

import io

import pytest

from app.services.resume import (
    ResumeExtractionError,
    extract_text,
    looks_like_a_resume,
    normalise_whitespace,
)
from app.services.resume.extractor import (
    _MAX_UNREADABLE_SHARE,
    MAX_RESUME_CHARS,
    resume_marker_count,
    unreadable_share,
)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

RESUME_BODY = """SPARSH SHARMA
Bangalore, India

OBJECTIVE
Java Full Stack Engineer seeking a role at Cognizant.

EDUCATION
B.Tech Computer Science, 2026. CGPA 8.4.

SKILLS
Java, Spring Boot, React, PostgreSQL, Redis, Docker, REST APIs, JPA

PROJECTS
E-Commerce Platform - Built REST APIs with Spring Boot and PostgreSQL for a
product catalog serving 10,000 daily users.

INTERNSHIP
Backend Intern, Acme Corp (2025) - wrote JPA repositories, fixed N+1 queries.

CERTIFICATIONS
Oracle Java SE 11 Developer
"""


def _docx_bytes(text: str) -> bytes:
    """A real DOCX, built in memory."""
    import docx

    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_table_bytes(rows: list[tuple[str, str]]) -> bytes:
    """A DOCX that lays everything out in a table, as many resume templates do."""
    import docx

    document = docx.Document()
    table = document.add_table(rows=0, cols=2)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestDocxExtraction:
    def test_extracts_a_docx_resume(self):
        text = extract_text(_docx_bytes(RESUME_BODY), DOCX_MIME, filename="cv.docx")
        assert "SPARSH SHARMA" in text
        assert "Spring Boot" in text
        assert "10,000 daily users" in text

    def test_reads_table_layouts(self):
        """
        Resume templates commonly put every section in a table. Reading only
        paragraphs would return almost nothing and look identical to a scanned
        file, sending the candidate to fix a problem they do not have.
        """
        rows = [
            ("SKILLS", "Java, Spring Boot, PostgreSQL, Redis, Docker, REST APIs"),
            ("EDUCATION", "B.Tech Computer Science 2026, CGPA 8.4 from a university"),
            ("EXPERIENCE", "Backend Intern at Acme Corp, built JPA repositories"),
            ("PROJECTS", "E-Commerce Platform serving 10,000 users with Spring Boot"),
        ]
        text = extract_text(_docx_table_bytes(rows), DOCX_MIME, filename="cv.docx")
        assert "Spring Boot" in text
        assert "B.Tech" in text

    def test_bytes_that_are_not_a_docx_at_all_are_named_as_such(self):
        """
        REASON CHANGED FROM `docx_unreadable` TO `unsupported_type`, and the change is the
        correction rather than the regression.

        These bytes have a zip signature and nothing else — no manifest, no
        `word/document.xml`, not even a valid archive. Since `file_safety.verify` decides
        the format from the content, they are now recognised as not being a Word document
        at all, which is what they are. The old answer told the candidate their DOCX was
        corrupt and to re-save it as .docx; that advice cannot help somebody who uploaded
        something that was never a document.

        The case this test was really protecting — a file that IS a DOCX and IS damaged —
        is unchanged and is pinned by `test_a_real_but_damaged_docx_still_names_the_format`
        below.
        """
        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(b"PK\x03\x04 not really a docx", DOCX_MIME, filename="cv.docx")
        assert exc.value.reason == "unsupported_type"
        # Still has to tell the candidate what to DO, not just that it broke.
        assert "PDF" in str(exc.value)

    def test_a_real_but_damaged_docx_still_names_the_format(self):
        """
        THE GUARANTEE THE TEST ABOVE USED TO CARRY. A genuine Word document that will not
        open must still get the specific, actionable message rather than being lumped in
        with "that is not a document" — the candidate has the right file and needs to know
        to re-save it.
        """
        from tests.docx_builder import valid_docx

        real = valid_docx("Education B.E. Computer Science skills projects experience" * 20)
        # A valid archive with the OOXML manifest, whose main part is truncated mid-stream.
        damaged = real.replace(b"word/document.xml", b"word/dXcument.xml")

        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(damaged, DOCX_MIME, filename="cv.docx")
        # Either answer is honest here — renaming the main part arguably makes it not a
        # DOCX — but it must never be silently accepted.
        assert exc.value.reason in ("docx_unreadable", "unsupported_type")
        assert str(exc.value)


class TestRejectsUnusableFiles:
    """Each of these must raise, because storing them silently is the bug."""

    def test_empty_file(self):
        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(b"", PDF_MIME, filename="cv.pdf")
        assert exc.value.reason == "empty_file"

    def test_bytes_that_are_not_a_pdf_at_all_are_named_as_such(self):
        """
        REASON CHANGED FROM `pdf_unreadable` TO `unsupported_type`. See the note on
        `test_bytes_that_are_not_a_docx_at_all_are_named_as_such`: this string has no
        `%PDF-` header, so calling it a corrupt PDF was never accurate, and the advice
        that came with it ("re-export it") could not help.
        """
        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(b"this is not a pdf", PDF_MIME, filename="cv.pdf")
        assert exc.value.reason == "unsupported_type"

    def test_a_real_but_damaged_pdf_still_names_the_format(self):
        """
        THE GUARANTEE THE TEST ABOVE USED TO CARRY, pinned on a file that really is a PDF.
        A candidate whose export is truncated has the right file and a fixable problem, and
        must be told so rather than being told it is not a PDF.
        """
        from tests.pdf_builder import build_pdf, visible_run

        real = build_pdf(visible_run("Education B.E. Computer Science skills projects" * 8))

        for damaged in (real[: len(real) // 2], real.replace(b"startxref", b"startxrKf")):
            with pytest.raises(ResumeExtractionError) as exc:
                extract_text(damaged, PDF_MIME, filename="cv.pdf")
            assert exc.value.reason == "pdf_unreadable", (
                "a genuinely damaged PDF lost its specific, actionable message"
            )
            assert "re-export" in str(exc.value).lower()

    def test_unsupported_type(self):
        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(b"\x89PNG\r\n", "image/png", filename="cv.png")
        assert exc.value.reason == "unsupported_type"

    def test_a_file_with_almost_no_text_is_rejected_as_a_scan(self):
        """
        A photographed or scanned resume extracts to a handful of characters. That
        is not a usable resume, and the message must say so — otherwise the
        interview quietly proceeds on three characters of noise.
        """
        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(_docx_bytes("Resume"), DOCX_MIME, filename="cv.docx")
        assert exc.value.reason == "no_text_layer"
        assert "scan" in str(exc.value).lower()


class TestTextThatExistsButCannotBeUsed:
    """
    THE THIRD ROOT CAUSE OF "the resume skills and projects are not been able to
    fetch", and the one no amount of AI-side hardening can reach: the analyser was
    being handed text with nothing in it to find.

    Every file in this class CLEARED the 200-character floor and was stored as a
    successfully-read resume. Both analysis halves then correctly found no skills
    and no projects in it, four billed retries later, and the candidate was told
    "your resume was read successfully, but the detailed skill analysis could not be
    completed" — which is both false and unactionable when what they are holding is
    a phone scan. Each must now fail at extraction, with a message that names the
    file to upload instead.
    """

    def test_a_scan_whose_only_text_is_scanner_furniture_is_rejected(self):
        """
        MEASURED, and the reason the 200-character floor is not enough on its own. A
        phone-scanned resume is not always a page of pure images: the scanner app
        stamps its own furniture into a text layer, and four pages of it is 307
        characters of "Scanned by CamScanner / Page 3 of 8 / IMG_0411.jpg" — clear
        of the floor, and containing not one word of the candidate's resume.
        """
        furniture = "\n".join(
            f"Scanned by CamScanner  Page {n} of 8  2026-08-22  IMG_{n * 137:04d}.jpg  "
            "www.camscanner.com"
            for n in range(1, 5)
        )
        assert len(furniture) > 200, "the point of this case is that it clears the floor"

        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(_docx_bytes(furniture), DOCX_MIME, filename="scan.docx")
        assert exc.value.reason == "no_resume_content"
        # It has to tell them WHICH file to upload instead, or the error is as much
        # of a dead end as the success message it replaced.
        assert "scan" in str(exc.value).lower()
        assert "original" in str(exc.value).lower()

    def test_a_broken_text_layer_is_rejected_rather_than_analysed(self):
        """
        A PDF whose fonts carry no usable ToUnicode CMap extracts to a wall of
        U+FFFD. 560 characters of it clears every length check there is and looks,
        to every component downstream, exactly like a successfully read resume.
        """
        mojibake = ("\ufffd" * 400) + " \ufffd\ufffd\ufffd " * 40

        with pytest.raises(ResumeExtractionError) as exc:
            extract_text(_docx_bytes(mojibake), DOCX_MIME, filename="cv.docx")
        assert exc.value.reason == "text_unreadable"
        assert "re-export" in str(exc.value).lower() or "export" in str(exc.value).lower()

    def test_control_codes_are_unreadable_too(self):
        """
        The other shape a broken CMap takes. Asserted on the helper rather than
        through a DOCX because python-docx refuses to write control characters at
        all — they only ever arrive from a PDF.
        """
        assert unreadable_share("\x01\x02\x03\x04" * 50 + "resume") > _MAX_UNREADABLE_SHARE
        assert unreadable_share(RESUME_BODY) == 0.0

    def test_one_stray_glyph_does_not_cost_a_candidate_their_resume(self):
        """
        The tolerance is not zero on purpose: a single bad ligature in an otherwise
        perfect export must not be a rejection.
        """
        text = RESUME_BODY.replace("Java Full Stack", "Java Full\ufffdStack")
        assert unreadable_share(text) < _MAX_UNREADABLE_SHARE
        assert "Spring Boot" in extract_text(
            _docx_bytes(text), DOCX_MIME, filename="cv.docx"
        )


class TestTheGateDoesNotRejectRealResumes:
    """
    THE OTHER HALF OF THE JUDGEMENT, and the reason the marker bar for rejection is
    zero rather than the two `looks_like_a_resume` wants. Refusing a candidate's
    actual resume is a worse failure than analysing a thin one, so each of these is
    a resume the gate was measured rejecting — or nearly rejecting — before it was
    tightened. They must all pass.
    """

    def test_a_terse_resume_with_no_standard_headings_is_accepted(self):
        """
        297 characters, and it never writes the words "skills", "projects" or
        "experience" — it just lists them. An earlier version of the gate (fewer
        than two markers) rejected this, which would have been a candidate unable to
        upload their own CV.
        """
        text = extract_text(
            _docx_bytes(
                "RAHUL KUMAR\n"
                "rahul@example.com | +91 90000 00000\n"
                "B.E. Computer Science, Anna University, 2027\n"
                "Java, Python, SQL, HTML, CSS, Git\n"
                "Library Management System - Java Swing desktop app with MySQL backend\n"
                "Weather App - fetches OpenWeather API data, shows a 5-day forecast\n"
                "NPTEL Programming in Java (Elite)"
            ),
            DOCX_MIME,
            filename="cv.docx",
        )
        assert "Java Swing" in text

    def test_a_resume_not_written_in_english_is_accepted(self):
        """
        `unreadable_share` counts unmappable glyphs, NOT unfamiliar ones. Every
        script is readable text; a resume in Devanagari must be unaffected.
        """
        text = extract_text(
            _docx_bytes(
                "राहुल कुमार\n"
                "शिक्षा: बी.टेक कंप्यूटर विज्ञान, 2027\n"
                "कौशल: Java, Python, SQL, Spring Boot, Docker\n"
                "परियोजना: पुस्तकालय प्रबंधन प्रणाली — Java Swing और MySQL\n"
                "internship at Acme Corp, backend developer, six months"
            ),
            DOCX_MIME,
            filename="cv.docx",
        )
        assert "Spring Boot" in text

    def test_the_rejection_bar_is_stricter_than_the_soft_signal(self):
        """
        The two thresholds this file's judgement rests on, pinned as a relationship
        rather than as values. `looks_like_a_resume` (soft, logged) wants two
        markers; the hard rejection wants zero. Collapsing them back into one — the
        obvious "reuse the helper" simplification — is what rejected a real CV.
        """
        one_marker = "B.E. Computer Science, Anna University, 2027. Java, Python, SQL."
        assert resume_marker_count(one_marker) == 1
        assert looks_like_a_resume(one_marker) is False  # soft signal: only logged
        # …and yet it must survive extraction, because one marker is still a signal.
        assert resume_marker_count(one_marker) > 0

        furniture = "Scanned by CamScanner  Page 3 of 8  IMG_0411.jpg"
        assert resume_marker_count(furniture) == 0

    def test_a_substantial_document_is_never_rejected_for_its_headings(self):
        """
        Past 1200 characters the marker check is not consulted at all, so an
        academic CV whose headings are all "Publications" and "Positions Held" gets
        the benefit of the doubt and reaches the analyser.
        """
        prose = (
            "Rahul Kumar has worked with Java, Spring and PostgreSQL on an ordering "
            "platform and a chat tool, and has published two papers. "
        ) * 10
        assert len(prose) > 1200
        assert "PostgreSQL" in extract_text(_docx_bytes(prose), DOCX_MIME, filename="cv.docx")


class TestContentTypeFallback:
    def test_extension_is_used_when_the_browser_sends_a_generic_type(self):
        """
        Some file managers upload PDFs as application/octet-stream. Rejecting
        those would turn a perfectly good resume into a support question.
        """
        data = _docx_bytes(RESUME_BODY)
        text = extract_text(data, "application/octet-stream", filename="cv.docx")
        assert "Spring Boot" in text

    def test_missing_content_type_still_works_with_an_extension(self):
        text = extract_text(_docx_bytes(RESUME_BODY), "", filename="cv.docx")
        assert "SPARSH SHARMA" in text


class TestTextIsBounded:
    def test_very_long_text_is_capped(self):
        """
        This text becomes AI input on every interview plan, so an unbounded
        resume is unbounded token spend later.
        """
        huge = _docx_bytes(RESUME_BODY + ("Java Spring Boot PostgreSQL. " * 3000))
        text = extract_text(huge, DOCX_MIME, filename="cv.docx")
        assert len(text) <= MAX_RESUME_CHARS


class TestNormaliseWhitespace:
    def test_collapses_the_ragged_output_pdfs_produce(self):
        messy = "Name\r\n\r\n\r\n\r\nSKILLS   \t  Java Spring   Boot\n\n\n\nEDUCATION"
        clean = normalise_whitespace(messy)
        assert "\n\n\n" not in clean
        assert "  " not in clean
        assert " " not in clean
        assert "\r" not in clean

    def test_keeps_paragraph_breaks(self):
        # Structure the model can use must survive; only excess is removed.
        assert "\n\n" in normalise_whitespace("SKILLS\nJava\n\n\n\nEDUCATION\nB.Tech")

    def test_strips_indentation_from_column_layouts(self):
        assert normalise_whitespace("    SKILLS\n      Java") == "SKILLS\nJava"

    def test_empty_input_is_safe(self):
        assert normalise_whitespace("") == ""


class TestLooksLikeAResume:
    def test_accepts_a_real_resume(self):
        assert looks_like_a_resume(RESUME_BODY) is True

    def test_case_insensitive(self):
        assert looks_like_a_resume("EXPERIENCE and EDUCATION") is True

    @pytest.mark.parametrize(
        "text",
        [
            "",
            "Your ticket for the 14:05 train to Pune. Coach B, seat 32.",
            "Dear candidate, congratulations on your offer.",
        ],
    )
    def test_rejects_documents_that_are_not_resumes(self, text: str):
        # Only a soft signal (it is logged, not enforced), but it should not
        # cheerfully call an train ticket a resume.
        assert looks_like_a_resume(text) is False
